"""Generate the 'Female Drivers Option' feasibility document as a .docx.

Run:  python _gen_female_drivers_doc.py
Output: Ziggo_Female_Drivers_Feature.docx  (in the repo root)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BRAND = RGBColor(0x1A, 0x3D, 0xB8)      # royal blue
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = DARK


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND
    return h


def body(text, bold=False, italic=False, color=None, size=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ============================ TITLE ============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Ziggo Super App")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = BRAND

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Feature Feasibility & Implementation Plan\nFemale-Driver Preference for Female Riders")
r.font.size = Pt(15)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Prepared for the product team  ·  Engineering scoping document")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY

doc.add_paragraph()

# ============================ 1. SUMMARY ============================
heading("1. Executive Summary", 1)
body(
    "This document scopes the work required to add a “Female Drivers” safety "
    "feature: when a verified female customer requests a ride, parcel or delivery, "
    "she can choose to be matched only with female drivers. The goal is to improve "
    "safety and comfort for women travelling alone, a feature women-focused ride "
    "services in South Asia (and PickMe’s own “Pick Me Lady”) already offer."
)
body(
    "Bottom line on difficulty: this is a MEDIUM-effort feature, not a hard one. "
    "The app already has everything structurally needed — a driver-matching engine, "
    "a booking flow, an admin approval pipeline and a profile system. What is missing "
    "is a single piece of information the app does not currently store: the gender of "
    "users and drivers. Most of the work is adding that field end-to-end and inserting "
    "one extra filter into the matching engine.",
    bold=False,
)
body(
    "Estimated effort: roughly 5–8 working days for one full-stack developer to "
    "ship a solid first version (backend + app + admin), plus a few days of testing "
    "and driver-onboarding/verification policy work. There are no third-party API "
    "costs — the change lives entirely inside the existing Ziggo codebase.",
    bold=True,
)

# ============================ 2. CURRENT STATE ============================
heading("2. Current State — Why It Isn’t “Just a Toggle”", 1)
body(
    "We reviewed the codebase to understand exactly what exists today. Findings:"
)
bullet("there is currently no gender field stored anywhere — not on the user account, "
       "not on the customer profile, and not on the driver profile. This is the core gap.",
       bold_prefix="No gender data: ")
bullet("the driver-matching engine selects drivers using only three rules — the driver is "
       "online, the driver is approved, and the driver’s vehicle type matches the request. "
       "There is no place where a gender preference could be applied yet.",
       bold_prefix="Matching is gender-blind: ")
bullet("the customer chooses a service (bike / tuk / car / van / truck), but has no way to "
       "express a preference like “female driver only” on the booking.",
       bold_prefix="Booking has no preference flag: ")
bullet("drivers are approved by an admin and upload documents (NIC, license), so we already "
       "have a human-in-the-loop step where a driver’s gender can be confirmed against their "
       "NIC — important, because a safety feature must be trustworthy.",
       bold_prefix="Admin approval already exists: ")
body(
    "So the feature is not blocked by anything architectural. It is a matter of adding one "
    "new piece of data and threading it through the layers that already exist.",
    italic=True,
)

# ============================ 3. WHAT NEEDS TO CHANGE ============================
heading("3. What Needs to Change (Plain-English)", 1)
body("The feature breaks into five concrete pieces of work:")

numbered("Add a “gender” field to user and driver records so the system knows who is "
         "female. Collected at sign-up / driver registration.",
         bold_prefix="Store gender. ")
numbered("Let a female customer turn on “Female drivers only” when booking — a simple "
         "switch on the ride/parcel screen. The option only appears for verified female users.",
         bold_prefix="Add the rider option. ")
numbered("Teach the matching engine one extra rule: if the request asks for female drivers, "
         "only consider female drivers. Everything else about matching stays the same.",
         bold_prefix="Filter the match. ")
numbered("Show the admin team a way to see and verify each driver’s gender during approval "
         "(checked against the uploaded NIC), and to report on female-driver supply.",
         bold_prefix="Admin verification. ")
numbered("Handle the “no female driver available” case gracefully — tell the rider, and "
         "offer to wait or to allow any driver, rather than silently failing.",
         bold_prefix="Handle no-supply. ")

# ============================ 4. TECHNICAL BREAKDOWN ============================
heading("4. Technical Breakdown (for Engineering)", 1)

heading("4.1 Backend (FastAPI + SQLAlchemy)", 2)
bullet("Add a gender column to the user model (values: female / male / other / undisclosed). "
       "Add the same to the driver registration flow, and a verified_female boolean on the "
       "driver that the admin sets at approval time (so the preference can only match drivers "
       "whose gender was actually confirmed).")
bullet("Run a database migration (Alembic) to add the new columns.")
bullet("Add a female_driver_only flag to the booking so the request carries the preference.")
bullet("Update the matching service to accept a female_only parameter and add one filter: "
       "when set, restrict candidates to drivers flagged verified-female. This is a small, "
       "localized change in one file.")
bullet("Update the booking-create endpoint and request schema to pass the flag through, and "
       "emit a clear ‘no female drivers available’ event when none are in range.")

heading("4.2 Customer App (Flutter)", 2)
bullet("Collect gender during profile creation / onboarding (and allow editing in profile).")
bullet("On the ride / flash / delivery booking screen, show a “Female drivers only” toggle "
       "— visible only when the signed-in customer is female. Pass the flag with the booking.")
bullet("Add a friendly empty-state when no female driver is found: ‘No female drivers nearby "
       "right now — wait, or allow any driver?’")
bullet("Optionally surface a small badge on the driver card so the rider sees confirmation.")

heading("4.3 Driver App (Flutter)", 2)
bullet("Add gender to the driver registration form.")
bullet("Optionally let female drivers opt in to a ‘ladies-only rides’ preference of their own.")

heading("4.4 Admin Panel (Jinja2)", 2)
bullet("Show driver gender on the driver detail / approval page, and a control to confirm "
       "‘verified female’ after checking the NIC document.")
bullet("Add a filter / count so operations can monitor female-driver supply by area.")

# ============================ 5. EFFORT TABLE ============================
heading("5. Effort & Complexity Estimate", 1)
body("Complexity is rated Low / Medium / High per work item. “Days” assume one "
     "experienced full-stack developer.")

table(
    ["Work Item", "Layer", "Complexity", "Est. Effort"],
    [
        ["Add gender + verified-female fields", "Backend / DB", "Low", "0.5 day"],
        ["Database migration", "Backend", "Low", "0.25 day"],
        ["Female-only flag on booking + schema", "Backend", "Low", "0.5 day"],
        ["Matching engine filter", "Backend", "Low–Medium", "0.5 day"],
        ["No-supply handling + event", "Backend", "Medium", "0.5 day"],
        ["Collect gender at onboarding", "Customer app", "Low", "0.5 day"],
        ["Booking-screen toggle + flow", "Customer app", "Medium", "1 day"],
        ["No-driver empty state + UX", "Customer app", "Medium", "0.5 day"],
        ["Driver registration gender field", "Driver app", "Low", "0.5 day"],
        ["Admin verify + supply reporting", "Admin panel", "Medium", "1 day"],
        ["End-to-end testing", "All", "Medium", "1–2 days"],
    ],
    widths=[2.6, 1.3, 1.2, 1.1],
)
body("Total engineering: ~5–8 working days for a robust first version.", bold=True)

# ============================ 6. NON-TECHNICAL ============================
heading("6. Non-Technical Considerations (Important)", 1)
body(
    "A safety feature is only as good as the trust behind it. These are policy decisions, "
    "not code, but they must be settled before launch:"
)
bullet("Gender must be verified against the driver’s NIC during approval — self-declaration "
       "alone is not safe for a women-only feature. The admin approval step is where this happens.",
       bold_prefix="Verification: ")
bullet("Decide how the rider proves she is female to unlock the option (typically the same "
       "NIC-based verification, or simply self-selection gated to verified accounts).",
       bold_prefix="Eligibility: ")
bullet("Store gender as optional and handle it sensitively; follow Sri Lanka data-protection "
       "norms and let users keep it undisclosed if they are not using the feature.",
       bold_prefix="Privacy: ")
bullet("In many areas there will be few female drivers. The feature must degrade gracefully "
       "and the business should actively recruit female drivers for it to be useful.",
       bold_prefix="Supply: ")
bullet("Define what happens for women-only requests that span multi-stop, flash parcel, "
       "rental and food/market delivery — the same filter applies, but confirm the policy.",
       bold_prefix="Scope: ")

# ============================ 7. RISKS ============================
heading("7. Risks & Mitigations", 1)
table(
    ["Risk", "Mitigation"],
    [
        ["Too few female drivers → long waits / no match",
         "Graceful fallback (wait or allow any driver); recruit female drivers; show ETA honestly."],
        ["Fake gender claims by drivers",
         "Admin verifies gender against NIC before approval; only verified-female drivers match."],
        ["Privacy / data-protection concerns",
         "Gender optional, stored minimally, used only for matching; clear consent at sign-up."],
        ["Existing drivers have no gender on record",
         "Backfill via admin during next approval cycle; default to undisclosed until set."],
    ],
    widths=[2.9, 3.4],
)

# ============================ 8. PHASING ============================
heading("8. Suggested Phasing", 1)
numbered("Foundation — add gender fields, migration, admin verification UI. Collect data, no "
         "rider-facing change yet.", bold_prefix="Phase 1: ")
numbered("Core feature — booking toggle, matching filter, no-supply handling, for standard "
         "rides first.", bold_prefix="Phase 2: ")
numbered("Expand — extend to flash/parcel, rental and delivery; add supply reporting and "
         "female-driver recruitment dashboards.", bold_prefix="Phase 3: ")

# ============================ 9. CONCLUSION ============================
heading("9. Conclusion", 1)
body(
    "Adding a female-driver option for female riders is a worthwhile and achievable feature. "
    "Technically it is a medium-effort change — the heavy lifting is a single new piece of "
    "data (gender) carried through the existing layers, plus one extra rule in the matching "
    "engine. The harder part is operational: verifying driver gender reliably and ensuring "
    "enough female drivers are available for the feature to feel real. With one developer it "
    "is realistically a one-to-two-week effort for a polished first release, requiring no new "
    "third-party services or costs."
)

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("— End of document —")
r.italic = True
r.font.color.rgb = GREY

out = "Ziggo_Female_Drivers_Feature.docx"
doc.save(out)
print("Wrote", out)
